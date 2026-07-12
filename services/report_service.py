from __future__ import annotations

import io
import json
import os
import time
from dataclasses import dataclass
from datetime import datetime
from typing import Any, Callable

import fitz
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

from config.api_key_store import load_google_api_key
from config.settings import GEMINI_SUMMARY_MODEL_DEFAULT
from services import CacheService, DocumentStorageService, LlmService, PdfService, TemplateService
from utils.docx_export import (
    COMPARISON_FONT_ANALYSIS,
    COMPARISON_FONT_OCR,
    COMPARISON_IMAGE_TARGET_DPI,
    add_comparison_page_banner,
    add_comparison_paragraph,
    add_comparison_text_body,
    append_docx_page_break,
    comparison_image_raster_plan,
    configure_document_landscape,
    landscape_content_size_inches,
)
from utils.reporting import (
    build_report_entry,
    discover_pdf_files,
    extract_jacar_ref_from_path,
    extract_transcription_text,
    load_index,
    now_iso,
    parse_analysis_page_json,
    safe_filename,
    write_json,
)

ProgressCallback = Callable[[int, int, str], None]


@dataclass
class ExportResult:
    success: int
    failed: int
    skipped: int
    manifest_path: str
    rows: list[dict[str, Any]]


class ReportService:
    def __init__(self, project_root: str) -> None:
        self.project_root = os.path.abspath(project_root)
        self.cache_service = CacheService()
        self.pdf_service = PdfService()
        self.storage_service = DocumentStorageService(
            project_root=self.project_root,
            cache_service=self.cache_service,
        )

    def default_index_path(self) -> str:
        return os.path.join(self.project_root, "Reports", "report_index.json")

    def default_comparison_dir(self) -> str:
        return os.path.join(self.project_root, "Reports", "comparison_docx")

    def default_summary_dir(self) -> str:
        return os.path.join(self.project_root, "Reports", "summary")

    @staticmethod
    def _manifest_name(prefix: str) -> str:
        ts = datetime.now().strftime("%Y%m%d_%H%M%S")
        return f"run_manifest_{prefix}_{ts}.json"

    @staticmethod
    def expected_comparison_docx_path(pdf_path: str, comparison_dir: str) -> str:
        """与 export_comparison_docx 输出命名规则一致。"""
        doc_name = os.path.basename(pdf_path)
        filename = safe_filename(os.path.splitext(doc_name)[0]) + "_按页对照报告.docx"
        return os.path.join(comparison_dir, filename)

    @staticmethod
    def expected_summary_md_path(pdf_path: str, summary_dir: str) -> str:
        """与 generate_summaries 输出命名规则一致。"""
        doc_name = os.path.basename(pdf_path)
        filename = safe_filename(os.path.splitext(doc_name)[0]) + ".summary.md"
        return os.path.join(summary_dir, filename)

    @staticmethod
    def _find_existing_export_file(
        output_dir: str,
        pdf_path: str,
        *,
        exact_path: str,
        ref: str,
        suffix: str,
    ) -> tuple[bool, str]:
        """
        在导出目录中定位已生成的汇报文件：先精确路径，再按 JACAR 编号模糊匹配。
        返回 (是否存在, 实际命中路径或预期路径)。
        """
        if exact_path and os.path.isfile(exact_path):
            return True, os.path.abspath(exact_path)

        if not ref or not os.path.isdir(output_dir):
            return False, exact_path

        suffix_lower = suffix.lower()
        ref_upper = ref.upper()
        candidates: list[str] = []
        for name in os.listdir(output_dir):
            if not name.lower().endswith(suffix_lower):
                continue
            if ref_upper in name.upper():
                candidates.append(os.path.join(output_dir, name))
        if not candidates:
            return False, exact_path
        candidates.sort(key=lambda p: (len(os.path.basename(p)), p))
        hit = os.path.abspath(candidates[0])
        return True, hit

    def probe_export_artifacts(
        self,
        pdf_path: str,
        *,
        comparison_dir: str | None = None,
        summary_dir: str | None = None,
    ) -> dict[str, Any]:
        """检查内容1 DOCX / 内容2 总结 MD 是否已在默认输出目录生成。"""
        comparison_dir = os.path.abspath(comparison_dir or self.default_comparison_dir())
        summary_dir = os.path.abspath(summary_dir or self.default_summary_dir())
        ref = extract_jacar_ref_from_path(pdf_path)
        bundle = self.storage_service.resolve_bundle_from_pdf(pdf_path)

        docx_expected = (
            self.storage_service.artifact_path(bundle, "export_comparison_docx")
            if bundle.layout == "bundle_v1"
            else self.expected_comparison_docx_path(pdf_path, comparison_dir)
        )
        docx_exists, docx_hit = self._find_existing_export_file(
            comparison_dir,
            pdf_path,
            exact_path=docx_expected,
            ref=ref,
            suffix=".docx",
        )

        md_expected = (
            self.storage_service.artifact_path(bundle, "summary")
            if bundle.layout == "bundle_v1"
            else self.expected_summary_md_path(pdf_path, summary_dir)
        )
        md_exists, md_hit = self._find_existing_export_file(
            summary_dir,
            pdf_path,
            exact_path=md_expected,
            ref=ref,
            suffix=".md",
        )

        notes: list[str] = []
        if docx_exists:
            notes.append(f"已生成内容1 DOCX：{os.path.basename(docx_hit)}")
        else:
            notes.append("未找到内容1 DOCX（Reports/comparison_docx）")
        if md_exists:
            notes.append(f"已生成内容2 总结 MD：{os.path.basename(md_hit)}")
        else:
            notes.append("未找到内容2 总结 MD（Reports/summary）")

        return {
            "comparison_docx_exists": docx_exists,
            "comparison_docx_path": docx_hit,
            "summary_md_exists": md_exists,
            "summary_md_path": md_hit,
            "export_notes": notes,
        }

    def build_index(
        self,
        *,
        pdf_root: str | None = None,
        output_path: str | None = None,
        comparison_dir: str | None = None,
        summary_dir: str | None = None,
    ) -> dict[str, Any]:
        pdf_root = os.path.abspath(pdf_root or os.path.join(self.project_root, "JACAR_Downloads"))
        output_path = os.path.abspath(output_path or self.default_index_path())
        comparison_dir = os.path.abspath(comparison_dir or self.default_comparison_dir())
        summary_dir = os.path.abspath(summary_dir or self.default_summary_dir())
        pdf_paths = discover_pdf_files(pdf_root)
        entries: list[dict[str, Any]] = []
        ready_count = 0
        docx_hit_count = 0
        md_hit_count = 0
        for pdf_path in pdf_paths:
            entry = build_report_entry(
                project_root=self.project_root,
                pdf_root=pdf_root,
                pdf_path=pdf_path,
                cache_service=self.cache_service,
                pdf_service=self.pdf_service,
            )
            if entry.ready:
                ready_count += 1
            row = entry.to_dict()
            export_info = self.probe_export_artifacts(
                pdf_path,
                comparison_dir=comparison_dir,
                summary_dir=summary_dir,
            )
            row.update(export_info)
            if export_info.get("comparison_docx_exists"):
                docx_hit_count += 1
            if export_info.get("summary_md_exists"):
                md_hit_count += 1
            entries.append(row)
        payload = {
            "generated_at": now_iso(),
            "project_root": self.project_root,
            "pdf_root": pdf_root,
            "comparison_dir": comparison_dir,
            "summary_dir": summary_dir,
            "total_documents": len(entries),
            "ready_documents": ready_count,
            "comparison_docx_found": docx_hit_count,
            "summary_md_found": md_hit_count,
            "entries": entries,
        }
        write_json(output_path, payload)
        payload["index_file"] = output_path
        return payload

    def load_index(self, index_path: str | None = None) -> dict[str, Any]:
        return load_index(os.path.abspath(index_path or self.default_index_path()))

    def delete_index(self, index_path: str | None = None) -> bool:
        target = os.path.abspath(index_path or self.default_index_path())
        if not os.path.exists(target):
            return False
        os.remove(target)
        return True

    @staticmethod
    def _analysis_lines(parsed: dict[str, Any]) -> list[str]:
        ctx = parsed.get("Historical_Context", {}) or {}
        entities = parsed.get("Entities_and_Concepts", {}) or {}
        disc = parsed.get("Discourse_Analysis", {}) or {}
        lines = [
            f"- Date_Written: {ctx.get('Date_Written', '') or '未知'}",
            f"- Author_Sender: {ctx.get('Author_Sender', '') or '未知'}",
            f"- Recipient: {ctx.get('Recipient', '') or '未知'}",
            f"- Document_Type: {ctx.get('Document_Type', '') or '未知'}",
            f"- Observation_Info: {disc.get('Observation_Info', '') or '未提及'}",
            f"- Core_Judgment: {disc.get('Core_Judgment', '') or '未提及'}",
            f"- Response_Action: {disc.get('Response_Action', '') or '未提及'}",
            f"- Relevance_Score: {disc.get('Relevance_Score', '') or '未设定'}",
        ]
        orgs = entities.get("Organizations", []) or []
        keys = entities.get("Key_Figures", []) or []
        if isinstance(orgs, list) and orgs:
            lines.append(f"- Organizations: {', '.join(str(x) for x in orgs[:12])}")
        if isinstance(keys, list) and keys:
            lines.append(f"- Key_Figures: {', '.join(str(x) for x in keys[:12])}")
        return lines

    @staticmethod
    def _write_page_block(
        doc: Document,
        *,
        pdf_page,
        page_index: int,
        total_pages: int,
        ocr_text: str,
        analysis_raw: str,
        include_analysis_raw: bool,
        image_dpi: float,
    ) -> None:
        """单页史料 → 三个横向 Word 页：扫描图满幅 / OCR 文本(MS Mincho) / 分析(宋体)。"""
        content_w_in, content_h_in = landscape_content_size_inches(doc)
        header_reserve_in = 0.55
        image_max_h_in = max(1.0, content_h_in - header_reserve_in - 0.25)
        image_max_w_in = max(1.0, content_w_in)
        rect = pdf_page.rect
        _fit_zoom, render_zoom, display_w_in = comparison_image_raster_plan(
            page_width_pt=rect.width,
            page_height_pt=rect.height,
            content_max_w_in=image_max_w_in,
            content_max_h_in=image_max_h_in,
            target_dpi=image_dpi,
        )

        # 1) OCR 原件（高 DPI 光栅化，Word 中按版面宽度显示）
        add_comparison_page_banner(doc, page_index=page_index, total_pages=total_pages)
        pix = pdf_page.get_pixmap(matrix=fitz.Matrix(render_zoom, render_zoom), alpha=False)
        img_buf = io.BytesIO(pix.tobytes("png"))
        image_paragraph = doc.add_paragraph()
        image_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        image_paragraph.paragraph_format.space_after = Pt(0)
        image_paragraph.add_run().add_picture(img_buf, width=Inches(display_w_in))
        append_docx_page_break(doc)

        # 2) OCR 提取文字
        add_comparison_page_banner(doc, page_index=page_index, total_pages=total_pages)
        add_comparison_text_body(doc, ocr_text, COMPARISON_FONT_OCR)
        append_docx_page_break(doc)

        # 3) 分析内容
        add_comparison_page_banner(doc, page_index=page_index, total_pages=total_pages)
        parsed, cleaned = parse_analysis_page_json(analysis_raw)
        if parsed is not None:
            for line in ReportService._analysis_lines(parsed):
                add_comparison_paragraph(doc, line, COMPARISON_FONT_ANALYSIS)
        else:
            add_comparison_text_body(doc, cleaned, COMPARISON_FONT_ANALYSIS)

        if include_analysis_raw and cleaned:
            add_comparison_paragraph(doc, "Analysis 原文：", COMPARISON_FONT_ANALYSIS)
            add_comparison_text_body(doc, cleaned, COMPARISON_FONT_ANALYSIS)

    def export_comparison_docx(
        self,
        *,
        index_data: dict[str, Any],
        selected_pdf_paths: set[str] | None,
        output_dir: str | None = None,
        include_incomplete: bool = False,
        include_analysis_raw: bool = False,
        image_dpi: float = COMPARISON_IMAGE_TARGET_DPI,
        progress_cb: ProgressCallback | None = None,
    ) -> ExportResult:
        use_bundle_output = output_dir is None and self.storage_service.is_bundle_layout()
        output_dir = os.path.abspath(output_dir or self.default_comparison_dir())
        os.makedirs(output_dir, exist_ok=True)
        entries = index_data.get("entries", []) or []
        selected = []
        for e in entries:
            pdf_path = str(e.get("pdf_path") or "")
            if selected_pdf_paths and pdf_path not in selected_pdf_paths:
                continue
            if not include_incomplete and not bool(e.get("ready", False)):
                continue
            selected.append(e)

        rows: list[dict[str, Any]] = []
        success = 0
        failed = 0
        total = len(selected)
        for i, entry in enumerate(selected, start=1):
            if progress_cb:
                progress_cb(i, total, f"导出内容1：{entry.get('pdf_rel_path')}")
            row: dict[str, Any] = {
                "pdf_rel_path": entry.get("pdf_rel_path"),
                "ready": bool(entry.get("ready")),
                "issues": entry.get("issues", []),
            }
            try:
                out = self._render_single_docx(
                    entry=entry,
                    output_dir=output_dir,
                    use_bundle_output=use_bundle_output,
                    include_analysis_raw=include_analysis_raw,
                    image_dpi=float(image_dpi),
                )
                row["status"] = "ok"
                row["output_docx"] = out
                success += 1
            except Exception as e:
                row["status"] = "failed"
                row["error"] = str(e)
                failed += 1
            rows.append(row)

        manifest = {
            "generated_at": now_iso(),
            "project_root": self.project_root,
            "output_dir": output_dir,
            "selected_documents": total,
            "success_documents": success,
            "failed_documents": failed,
            "rows": rows,
        }
        manifest_path = os.path.join(output_dir, self._manifest_name("comparison"))
        write_json(manifest_path, manifest)
        return ExportResult(success=success, failed=failed, skipped=0, manifest_path=manifest_path, rows=rows)

    def _render_single_docx(
        self,
        *,
        entry: dict[str, Any],
        output_dir: str,
        use_bundle_output: bool,
        include_analysis_raw: bool,
        image_dpi: float,
    ) -> str:
        pdf_path = entry["pdf_path"]
        page_count = int(entry.get("page_count", 0))
        ocr_pages = self.cache_service.read_paged_cache(entry["ocr_cache_path"])
        analysis_pages = self.cache_service.read_paged_cache(entry["analysis_cache_path"])

        doc = Document()
        configure_document_landscape(doc)
        doc.add_heading(f"史料按页对照报告：{os.path.basename(pdf_path)}", level=1)
        doc.add_paragraph(f"来源文件：{entry.get('pdf_rel_path', pdf_path)}")
        doc.add_paragraph(f"总页数：{page_count}")

        with fitz.open(pdf_path) as pdf:
            total = min(page_count, len(pdf))
            for i in range(total):
                ocr_text = extract_transcription_text(ocr_pages[i] if i < len(ocr_pages) else "")
                analysis_raw = analysis_pages[i] if i < len(analysis_pages) else ""
                self._write_page_block(
                    doc,
                    pdf_page=pdf[i],
                    page_index=i,
                    total_pages=total,
                    ocr_text=ocr_text,
                    analysis_raw=analysis_raw,
                    include_analysis_raw=include_analysis_raw,
                    image_dpi=image_dpi,
                )
                if i < total - 1:
                    append_docx_page_break(doc)

        if use_bundle_output:
            bundle = self.storage_service.resolve_bundle_from_pdf(pdf_path)
            out_path = self.storage_service.resolve_write_path(bundle, "export_comparison_docx")
        else:
            name = safe_filename(os.path.splitext(os.path.basename(pdf_path))[0]) + "_按页对照报告.docx"
            out_path = os.path.join(output_dir, name)
        os.makedirs(os.path.dirname(os.path.abspath(out_path)), exist_ok=True)
        doc.save(out_path)
        return out_path

    @staticmethod
    def _render_summary_prompt(*, doc_name: str, page_count: int) -> str:
        try:
            return TemplateService().render_prompt(
                "report_summary_prompt.jinja",
                {"doc_name": doc_name, "page_count": page_count},
            )
        except Exception:
            return (
                f"请基于以下 Analysis 数据，生成《{doc_name}》的研究汇报摘要。\n"
                "要求：概述背景时间线、观察判断因应主线、高频组织人物、研究价值与证据摘录。"
            )

    @staticmethod
    def _render_partial_summary_prompt(
        *,
        doc_name: str,
        page_count: int,
        chunk_index: int,
        chunk_total: int,
        page_start: int,
        page_end: int,
    ) -> str:
        return (
            f"你将收到《{doc_name}》的 Analysis 分段数据（第 {chunk_index}/{chunk_total} 批，"
            f"对应原文第 {page_start}-{page_end} 页，共 {page_count} 页中的一部分）。\n"
            "请输出该批次的小结（Markdown）：\n"
            "1) 该批次核心观察与判断；\n"
            "2) 关键组织/人物/关键词；\n"
            "3) 该批次的研究价值与证据页码。\n"
            "要求：只基于本批次内容，不要臆测全局结论。"
        )

    @staticmethod
    def _is_retryable_error(err: Exception) -> bool:
        msg = str(err).lower()
        keys = [
            "请求超时",
            "timed out",
            "timeout",
            "429",
            "rate limit",
            "503",
            "500",
            "connection reset",
            "temporarily unavailable",
            "network",
        ]
        return any(k in msg for k in keys)

    def _call_with_retry(self, fn: Callable[[], tuple[str, dict[str, Any]]], *, retry_attempts: int) -> tuple[str, dict[str, Any], int]:
        attempts = max(1, int(retry_attempts))
        last_err: Exception | None = None
        for attempt in range(1, attempts + 1):
            try:
                text, usage = fn()
                return text, usage, attempt
            except Exception as e:
                last_err = e
                if attempt >= attempts or not self._is_retryable_error(e):
                    raise
                wait_s = min(16, 2 ** attempt)
                time.sleep(wait_s)
        assert last_err is not None
        raise last_err

    @staticmethod
    def _build_analysis_chunks(
        analysis_pages: list[str],
        *,
        max_chars_per_chunk: int,
        chunk_pages: int,
    ) -> tuple[list[dict[str, Any]], bool]:
        units: list[tuple[int, str]] = []
        truncated = False
        for idx, page_raw in enumerate(analysis_pages, start=1):
            parsed, cleaned = parse_analysis_page_json(page_raw)
            if parsed is not None:
                ctx = parsed.get("Historical_Context", {}) or {}
                entities = parsed.get("Entities_and_Concepts", {}) or {}
                disc = parsed.get("Discourse_Analysis", {}) or {}
                obj = {
                    "page": idx,
                    "Date_Written": ctx.get("Date_Written"),
                    "Author_Sender": ctx.get("Author_Sender"),
                    "Recipient": ctx.get("Recipient"),
                    "Document_Type": ctx.get("Document_Type"),
                    "Organizations": (entities.get("Organizations", []) or [])[:12],
                    "Key_Figures": (entities.get("Key_Figures", []) or [])[:12],
                    "Discourse_Keywords": (entities.get("Discourse_Keywords", []) or [])[:20],
                    "Observation_Info": disc.get("Observation_Info"),
                    "Core_Judgment": disc.get("Core_Judgment"),
                    "Response_Action": disc.get("Response_Action"),
                    "Relevance_Score": disc.get("Relevance_Score"),
                }
                page_chunk = f"[PAGE {idx}]\n{json.dumps(obj, ensure_ascii=False, indent=2)}"
            else:
                page_chunk = f"[PAGE {idx}]\n{cleaned[:2500]}"
            units.append((idx, page_chunk))

        chunks: list[dict[str, Any]] = []
        current_units: list[str] = []
        current_pages: list[int] = []
        max_chars = max(5000, int(max_chars_per_chunk))
        max_pages = max(1, int(chunk_pages))

        def flush_current() -> None:
            if not current_units:
                return
            chunks.append(
                {
                    "text": "\n\n".join(current_units),
                    "page_start": current_pages[0],
                    "page_end": current_pages[-1],
                    "page_count": len(current_pages),
                }
            )
            current_units.clear()
            current_pages.clear()

        for page_no, unit in units:
            if len(unit) > max_chars:
                truncated = True
                unit = unit[:max_chars]
            should_flush = False
            if current_units and len(current_pages) >= max_pages:
                should_flush = True
            elif current_units and (len("\n\n".join(current_units)) + 2 + len(unit) > max_chars):
                should_flush = True
            if should_flush:
                flush_current()
            current_units.append(unit)
            current_pages.append(page_no)
        flush_current()
        return chunks, truncated

    def generate_summaries(
        self,
        *,
        index_data: dict[str, Any],
        selected_pdf_paths: set[str] | None,
        output_dir: str | None = None,
        api_key: str | None = None,
        model_name: str = GEMINI_SUMMARY_MODEL_DEFAULT,
        include_incomplete: bool = False,
        max_source_chars: int = 120000,
        chunk_pages: int = 20,
        retry_attempts: int = 3,
        skip_existing: bool = True,
        progress_cb: ProgressCallback | None = None,
    ) -> ExportResult:
        use_bundle_output = output_dir is None and self.storage_service.is_bundle_layout()
        output_dir = os.path.abspath(output_dir or self.default_summary_dir())
        os.makedirs(output_dir, exist_ok=True)
        entries = index_data.get("entries", []) or []
        selected = []
        for e in entries:
            pdf_path = str(e.get("pdf_path") or "")
            if selected_pdf_paths and pdf_path not in selected_pdf_paths:
                continue
            if not include_incomplete and not bool(e.get("ready", False)):
                continue
            selected.append(e)

        key = (
            (api_key or "").strip()
            or os.getenv("GOOGLE_GEMINI_API_KEY", "").strip()
            or os.getenv("GOOGLE_VISION_API_KEY", "").strip()
            or load_google_api_key()
        )
        if not key:
            raise RuntimeError("未找到 Gemini API Key。请先在设置页配置。")

        llm = LlmService(api_key=key, project_root=self.project_root, timeout_sec=180)
        rows: list[dict[str, Any]] = []
        success = 0
        failed = 0
        skipped = 0
        total = len(selected)
        for i, entry in enumerate(selected, start=1):
            if progress_cb:
                progress_cb(i, total, f"导出内容2：{entry.get('pdf_rel_path')}")
            row: dict[str, Any] = {
                "pdf_rel_path": entry.get("pdf_rel_path"),
                "ready": bool(entry.get("ready")),
                "issues": entry.get("issues", []),
                "model": model_name,
            }
            try:
                doc_name = os.path.basename(entry["pdf_path"])
                if use_bundle_output:
                    bundle = self.storage_service.resolve_bundle_from_pdf(entry["pdf_path"])
                    out_path = self.storage_service.resolve_write_path(bundle, "summary")
                else:
                    out_name = safe_filename(os.path.splitext(doc_name)[0]) + ".summary.md"
                    out_path = os.path.join(output_dir, out_name)
                if skip_existing and os.path.isfile(out_path):
                    row["status"] = "skipped_existing"
                    row["output_summary"] = out_path
                    row["skip_reason"] = "summary_exists"
                    skipped += 1
                    rows.append(row)
                    continue

                analysis_pages = self.cache_service.read_paged_cache(entry["analysis_cache_path"])
                page_count = int(entry.get("page_count", 0))
                chunks, source_truncated = self._build_analysis_chunks(
                    analysis_pages,
                    max_chars_per_chunk=max(5000, int(max_source_chars)),
                    chunk_pages=chunk_pages,
                )
                if not chunks:
                    raise RuntimeError("Analysis 缓存为空，无法生成总结。")

                partial_texts: list[str] = []
                partial_attempts: list[int] = []
                chunk_total = len(chunks)
                for ci, chunk in enumerate(chunks, start=1):
                    if progress_cb:
                        progress_cb(i, total, f"内容2分段总结：{entry.get('pdf_rel_path')} [{ci}/{chunk_total}]")
                    partial_prompt = self._render_partial_summary_prompt(
                        doc_name=doc_name,
                        page_count=page_count,
                        chunk_index=ci,
                        chunk_total=chunk_total,
                        page_start=int(chunk["page_start"]),
                        page_end=int(chunk["page_end"]),
                    )
                    partial_text, _partial_usage, attempt_used = self._call_with_retry(
                        lambda pp=partial_prompt, ct=chunk["text"]: llm.detect_text(
                            screen_name="ReportSummaryUI",
                            task_name="进度汇报总结-分段",
                            selected_pdf_path=entry["pdf_path"],
                            file_name=f"{doc_name}_summary_chunk_{ci}",
                            model_name=model_name,
                            academic_prompt=pp,
                            behavior_name="进度汇报总结-分段",
                            page_index=None,
                            source_text=ct,
                        ),
                        retry_attempts=retry_attempts,
                    )
                    partial_attempts.append(attempt_used)
                    partial_texts.append(
                        f"[CHUNK {ci} | p{chunk['page_start']}-{chunk['page_end']}]\n{partial_text.strip()}"
                    )

                final_source = "\n\n".join(partial_texts)
                final_prompt = self._render_summary_prompt(doc_name=doc_name, page_count=page_count)
                summary_text, usage_summary, final_attempt = self._call_with_retry(
                    lambda: llm.detect_text(
                        screen_name="ReportSummaryUI",
                        task_name="进度汇报总结",
                        selected_pdf_path=entry["pdf_path"],
                        file_name=f"{doc_name}_summary",
                        model_name=model_name,
                        academic_prompt=final_prompt,
                        behavior_name="进度汇报总结",
                        page_index=None,
                        source_text=final_source,
                    ),
                    retry_attempts=retry_attempts,
                )
                with open(out_path, "w", encoding="utf-8") as f:
                    f.write(f"# {doc_name}｜分析总结\n\n")
                    if source_truncated:
                        f.write("> 注：源数据超长，分段时存在截断。\n\n")
                    if chunk_total > 1:
                        f.write(f"> 注：本总结由 {chunk_total} 个分段小结二次汇总生成。\n\n")
                    f.write(summary_text.strip() + "\n")

                row["status"] = "ok"
                row["output_summary"] = out_path
                row["analysis_pages"] = len(analysis_pages)
                row["source_truncated"] = source_truncated
                row["chunk_count"] = chunk_total
                row["partial_attempts"] = partial_attempts
                row["final_attempt"] = final_attempt
                row["usage_summary"] = usage_summary
                success += 1
            except Exception as e:
                row["status"] = "failed"
                row["error"] = str(e)
                failed += 1
            rows.append(row)

        manifest = {
            "generated_at": now_iso(),
            "project_root": self.project_root,
            "output_dir": output_dir,
            "selected_documents": total,
            "success_documents": success,
            "failed_documents": failed,
            "skipped_documents": skipped,
            "rows": rows,
        }
        manifest_path = os.path.join(output_dir, self._manifest_name("summary"))
        write_json(manifest_path, manifest)
        return ExportResult(success=success, failed=failed, skipped=skipped, manifest_path=manifest_path, rows=rows)
