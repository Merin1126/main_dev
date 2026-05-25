from __future__ import annotations

import argparse
import io
import os
import sys
from pathlib import Path
from typing import Any

import fitz
from docx import Document
from docx.shared import Inches

PROJECT_ROOT = Path(__file__).resolve().parent.parent
if str(PROJECT_ROOT) not in sys.path:
    sys.path.insert(0, str(PROJECT_ROOT))

from services import CacheService
from utils.reporting import (
    extract_transcription_text,
    load_index,
    now_iso,
    parse_analysis_page_json,
    safe_filename,
    write_json,
)


def _default_index_path(project_root: str) -> str:
    return os.path.join(project_root, "Reports", "report_index.json")


def _default_output_dir(project_root: str) -> str:
    return os.path.join(project_root, "Reports", "comparison_docx")


def _default_manifest_path(output_dir: str) -> str:
    return os.path.join(output_dir, f"run_manifest_comparison_{now_iso().replace(':', '-')}.json")


def _analysis_lines(parsed: dict[str, Any]) -> list[str]:
    ctx = parsed.get("Historical_Context", {}) or {}
    entities = parsed.get("Entities_and_Concepts", {}) or {}
    disc = parsed.get("Discourse_Analysis", {}) or {}
    lines = [
        f"- Date_Written: {ctx.get('Date_Written', '') or '未知'}",
        f"- Author_Sender: {ctx.get('Author_Sender', '') or '未知'}",
        f"- Recipient: {ctx.get('Recipient', '') or '未知'}",
        f"- Document_Type: {ctx.get('Document_Type', '') or '未知'}",
        f"- Observation_Info: {disc.get('Observation_Info', '') or '未提及'}",
        f"- Core_Judgment: {disc.get('Core_Judgment', '') or '未提及'}",
        f"- Response_Action: {disc.get('Response_Action', '') or '未提及'}",
        f"- Relevance_Score: {disc.get('Relevance_Score', '') or '未设定'}",
    ]
    orgs = entities.get("Organizations", []) or []
    keys = entities.get("Key_Figures", []) or []
    if isinstance(orgs, list) and orgs:
        lines.append(f"- Organizations: {', '.join(str(x) for x in orgs[:12])}")
    if isinstance(keys, list) and keys:
        lines.append(f"- Key_Figures: {', '.join(str(x) for x in keys[:12])}")
    return lines


def _write_page_block(
    doc: Document,
    *,
    pdf_page,
    page_index: int,
    total_pages: int,
    ocr_text: str,
    analysis_raw: str,
    include_analysis_raw: bool,
    image_zoom: float,
) -> None:
    doc.add_heading(f"第 {page_index + 1} / {total_pages} 页", level=2)
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    table.autofit = False
    table.columns[0].width = Inches(3.0)
    table.columns[1].width = Inches(3.5)

    left = table.cell(0, 0)
    right = table.cell(0, 1)

    pix = pdf_page.get_pixmap(matrix=fitz.Matrix(image_zoom, image_zoom), alpha=False)
    img_buf = io.BytesIO(pix.tobytes("png"))
    p_img = left.paragraphs[0]
    p_img.add_run().add_picture(img_buf, width=Inches(2.9))

    right.paragraphs[0].add_run("OCR 提取：")
    right.add_paragraph(ocr_text or "（空）")

    doc.add_paragraph("Analysis：")
    parsed, cleaned = parse_analysis_page_json(analysis_raw)
    if parsed is not None:
        for line in _analysis_lines(parsed):
            doc.add_paragraph(line)
    else:
        doc.add_paragraph(cleaned or "（空）")

    if include_analysis_raw and cleaned:
        doc.add_paragraph("Analysis 原文：")
        doc.add_paragraph(cleaned)


def _render_single_docx(
    *,
    entry: dict[str, Any],
    cache_service: CacheService,
    output_dir: str,
    include_analysis_raw: bool,
    image_zoom: float,
) -> str:
    pdf_path = entry["pdf_path"]
    page_count = int(entry.get("page_count", 0))
    ocr_pages = cache_service.read_paged_cache(entry["ocr_cache_path"])
    analysis_pages = cache_service.read_paged_cache(entry["analysis_cache_path"])

    doc = Document()
    doc.add_heading(f"史料按页对照报告：{os.path.basename(pdf_path)}", level=1)
    doc.add_paragraph(f"来源文件：{entry.get('pdf_rel_path', pdf_path)}")
    doc.add_paragraph(f"总页数：{page_count}")

    with fitz.open(pdf_path) as pdf:
        total = min(page_count, len(pdf))
        for i in range(total):
            ocr_text = extract_transcription_text(ocr_pages[i] if i < len(ocr_pages) else "")
            analysis_raw = analysis_pages[i] if i < len(analysis_pages) else ""
            _write_page_block(
                doc,
                pdf_page=pdf[i],
                page_index=i,
                total_pages=total,
                ocr_text=ocr_text,
                analysis_raw=analysis_raw,
                include_analysis_raw=include_analysis_raw,
                image_zoom=image_zoom,
            )
            if i < total - 1:
                doc.add_page_break()

    name = safe_filename(os.path.splitext(os.path.basename(pdf_path))[0]) + "_按页对照报告.docx"
    out_path = os.path.join(output_dir, name)
    doc.save(out_path)
    return out_path


def main() -> int:
    parser = argparse.ArgumentParser(description="Export per-page comparison Word report.")
    parser.add_argument("--project-root", default=str(PROJECT_ROOT), help="Project root path.")
    parser.add_argument("--index-file", default=None, help="Index file path.")
    parser.add_argument("--output-dir", default=None, help="Output directory for docx files.")
    parser.add_argument("--manifest", default=None, help="Optional manifest output path.")
    parser.add_argument("--include-incomplete", action="store_true", help="Include incomplete documents.")
    parser.add_argument("--include-analysis-raw", action="store_true", help="Append raw analysis text/json.")
    parser.add_argument("--max-docs", type=int, default=0, help="Process up to N documents (0 = all).")
    parser.add_argument("--image-zoom", type=float, default=1.5, help="PDF render zoom for page snapshot.")
    args = parser.parse_args()

    project_root = os.path.abspath(args.project_root)
    index_file = os.path.abspath(args.index_file or _default_index_path(project_root))
    output_dir = os.path.abspath(args.output_dir or _default_output_dir(project_root))
    os.makedirs(output_dir, exist_ok=True)

    index_data = load_index(index_file)
    entries = index_data.get("entries", []) or []
    selected = [
        e
        for e in entries
        if bool(e.get("ready", False)) or args.include_incomplete
    ]
    if args.max_docs and args.max_docs > 0:
        selected = selected[: args.max_docs]

    cache_service = CacheService()
    manifest_rows: list[dict[str, Any]] = []
    success = 0
    failed = 0
    for entry in selected:
        row: dict[str, Any] = {
            "pdf_rel_path": entry.get("pdf_rel_path"),
            "ready": bool(entry.get("ready")),
            "issues": entry.get("issues", []),
        }
        try:
            out = _render_single_docx(
                entry=entry,
                cache_service=cache_service,
                output_dir=output_dir,
                include_analysis_raw=args.include_analysis_raw,
                image_zoom=max(0.8, float(args.image_zoom)),
            )
            row["status"] = "ok"
            row["output_docx"] = out
            success += 1
            print(f"[OK] {entry.get('pdf_rel_path')} -> {out}")
        except Exception as e:
            row["status"] = "failed"
            row["error"] = str(e)
            failed += 1
            print(f"[FAILED] {entry.get('pdf_rel_path')}: {e}")
        manifest_rows.append(row)

    manifest = {
        "generated_at": now_iso(),
        "project_root": project_root,
        "index_file": index_file,
        "output_dir": output_dir,
        "selected_documents": len(selected),
        "success_documents": success,
        "failed_documents": failed,
        "rows": manifest_rows,
    }
    manifest_path = os.path.abspath(args.manifest or _default_manifest_path(output_dir))
    write_json(manifest_path, manifest)
    print(f"Manifest written: {manifest_path}")
    print(f"Done. success={success}, failed={failed}")
    return 0 if failed == 0 else 2


if __name__ == "__main__":
    raise SystemExit(main())
