#!/usr/bin/env python3
"""Generate vertical historical-record DOCX files from paragraph-flow JSON.

Usage:
    python vertical_historical_docx_generator.py input.json output.docx

Dependency:
    pip install python-docx

The JSON format is demonstrated by:
    琿警機密第六四一號_段落自动换行示例.json
"""

import argparse
import json
import re
import sqlite3
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT, WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


DEFAULT_FONT = "MS Mincho"


def default_hrs_db_path(project_root: str | Path | None = None) -> Path:
    if project_root:
        return Path(project_root).expanduser().resolve() / "database" / "hrs.sqlite3"
    return Path(__file__).resolve().parents[2] / "database" / "hrs.sqlite3"
ALIGNMENT_MAP = {
    "left": WD_ALIGN_PARAGRAPH.LEFT,
    "center": WD_ALIGN_PARAGRAPH.CENTER,
    "right": WD_ALIGN_PARAGRAPH.RIGHT,
    "justify": WD_ALIGN_PARAGRAPH.JUSTIFY,
}
DEFAULT_ROLE_STYLES = {
    "header": {"first_line_indent_chars": 0},
    "date": {"first_line_indent_chars": 1},
    "sender": {"first_line_indent_chars": 2},
    "recipient": {"first_line_indent_chars": 2},
    "title": {
        "alignment": "center",
        "bold": True,
        "space_before_lines": 0.5,
        "space_after_lines": 0.5,
    },
}
NUMERAL_PATTERN = r"(?:[一二三四五六七八九十百千〇零壱弐参\d０-９]+)"
BRACKET_OPEN_CHARS = r"（(︵【［\[\{｛〔「『〈《〖〚〘"
BRACKET_CLOSE_CHARS = r"）)︶】］\]\}｝〕」』〉》〗〛〙"
NUMBERED_ITEM_RE = re.compile(
    rf"^\s*(?:(?:[{BRACKET_OPEN_CHARS}]\s*{NUMERAL_PATTERN}\s*[{BRACKET_CLOSE_CHARS}])|(?:{NUMERAL_PATTERN}[、．.。]))"
)


PUNCT_MAP = str.maketrans(
    {
        ",": "，",
        ".": "．",
        ":": "：",
        ";": "；",
        "!": "！",
        "?": "？",
        "[": "［",
        "]": "］",
        "{": "｛",
        "}": "｝",
        "<": "＜",
        ">": "＞",
        "/": "／",
        "\\": "＼",
        "-": "－",
        "~": "～",
        " ": "　",
    }
)

VERTICAL_BRACKET_MAP = str.maketrans(
    {
        "（": "︵",
        "）": "︶",
        "(": "︵",
        ")": "︶",
        "［": "﹇",
        "］": "﹈",
        "[": "﹇",
        "]": "﹈",
        "｛": "︷",
        "｝": "︸",
        "{": "︷",
        "}": "︸",
        "〔": "︹",
        "〕": "︺",
        "【": "︻",
        "】": "︼",
        "〖": "︗",
        "〗": "︘",
        "〘": "︗",
        "〙": "︘",
        "〚": "︗",
        "〛": "︘",
        "「": "﹁",
        "」": "﹂",
        "『": "﹃",
        "』": "﹄",
        "〈": "︿",
        "〉": "﹀",
        "《": "︽",
        "》": "︾",
    }
)


def set_page_border(section, border):
    if not border.get("enabled", True):
        return
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        node = pg_borders.find(qn(f"w:{edge}"))
        if node is None:
            node = OxmlElement(f"w:{edge}")
            pg_borders.append(node)
        node.set(qn("w:val"), "single")
        node.set(qn("w:sz"), str(border.get("size", 6)))
        node.set(qn("w:space"), str(border.get("space", 4)))
        node.set(qn("w:color"), border.get("color", "000000"))


def set_native_vertical_section(section, line_pitch):
    sect_pr = section._sectPr
    text_direction = sect_pr.find(qn("w:textDirection"))
    if text_direction is None:
        text_direction = OxmlElement("w:textDirection")
        sect_pr.append(text_direction)
    text_direction.set(qn("w:val"), "tbRl")

    doc_grid = sect_pr.find(qn("w:docGrid"))
    if doc_grid is None:
        doc_grid = OxmlElement("w:docGrid")
        sect_pr.append(doc_grid)
    doc_grid.set(qn("w:type"), "linesAndChars")
    doc_grid.set(qn("w:linePitch"), str(line_pitch))
    doc_grid.set(qn("w:charSpace"), "0")


def set_run_font(run, font, size, bold):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold


def set_paragraph_font(paragraph, font, size, bold=False):
    for run in paragraph.runs:
        set_run_font(run, font, size, bold)


def normalize_text(text, normalization):
    if normalization.get("fullwidth_ascii_digits", True) or normalization.get("fullwidth_ascii_letters", True):
        converted = []
        for ch in text:
            if normalization.get("fullwidth_ascii_digits", True) and "0" <= ch <= "9":
                converted.append(chr(ord(ch) + 0xFEE0))
            elif normalization.get("fullwidth_ascii_letters", True) and ("A" <= ch <= "Z" or "a" <= ch <= "z"):
                converted.append(chr(ord(ch) + 0xFEE0))
            else:
                converted.append(ch)
        text = "".join(converted)
    if normalization.get("vertical_parentheses", True):
        text = text.translate(VERTICAL_BRACKET_MAP)
    if normalization.get("fullwidth_punctuation", True):
        text = text.translate(PUNCT_MAP)
    if normalization.get("normalize_numbered_marker", True):
        text = re.sub(rf"^({NUMERAL_PATTERN}[、．.。])\s*", r"\1　", text)
    return text


def infer_role(block):
    role = block.get("role", "body")
    if role != "body":
        return role
    if NUMBERED_ITEM_RE.match(block.get("text", "")):
        return "numbered_item"
    return role


def get_role_style(layout, role):
    styles = dict(DEFAULT_ROLE_STYLES)
    styles.update(layout.get("role_styles", {}))
    return styles.get(role, {})


def get_block_style(block, layout, role, name, default=None):
    role_style = get_role_style(layout, role)
    return block.get(name, role_style.get(name, default))


def apply_section_layout(section, layout):
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width = Cm(layout["page"]["width_cm"])
    section.page_height = Cm(layout["page"]["height_cm"])
    section.top_margin = Cm(layout["page"]["margins_cm"]["top"])
    section.bottom_margin = Cm(layout["page"]["margins_cm"]["bottom"])
    section.left_margin = Cm(layout["page"]["margins_cm"]["left"])
    section.right_margin = Cm(layout["page"]["margins_cm"]["right"])
    section.header_distance = Cm(0.2)
    section.footer_distance = Cm(layout.get("folio", {}).get("footer_distance_cm", 0.2))
    set_page_border(section, layout.get("border", {}))
    set_native_vertical_section(section, layout.get("line_pitch", 220))


def apply_layout(doc, layout):
    apply_section_layout(doc.sections[0], layout)
    normal = doc.styles["Normal"]
    normal.font.name = layout.get("font", DEFAULT_FONT)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), layout.get("font", DEFAULT_FONT))
    normal.font.size = Pt(layout.get("body_size_pt", 11.8))
    normal.paragraph_format.space_after = Pt(0)
    normal.paragraph_format.line_spacing = 1.0


def add_paragraph_block(doc, block, layout):
    if block.get("role") == "page_number":
        return

    size = block.get("size_pt", layout.get("body_size_pt", 11.8))
    role = infer_role(block)
    role_style = get_role_style(layout, role)
    p = doc.add_paragraph()
    alignment = block.get("alignment", role_style.get("alignment"))
    if alignment:
        p.alignment = ALIGNMENT_MAP.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    pf = p.paragraph_format
    space_before = get_block_style(block, layout, role, "space_before_pt")
    if space_before is None:
        space_before = size * get_block_style(block, layout, role, "space_before_lines", 0)
    space_after = get_block_style(block, layout, role, "space_after_pt")
    if space_after is None:
        space_after = size * get_block_style(block, layout, role, "space_after_lines", 0)
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = get_block_style(block, layout, role, "line_spacing", 1.0)

    first_line_indent_chars = get_block_style(block, layout, role, "first_line_indent_chars", 0)
    if role == "numbered_item":
        hanging_chars = block.get(
            "hanging_indent_chars",
            layout.get("numbered_item", {}).get("hanging_indent_chars", 2),
        )
        pf.left_indent = Pt(size * hanging_chars)
        pf.first_line_indent = Pt(-size * hanging_chars)
    elif first_line_indent_chars:
        pf.first_line_indent = Pt(size * first_line_indent_chars)

    text = normalize_text(block["text"], layout.get("normalization", {}))
    run = p.add_run(text)
    set_run_font(
        run,
        font=block.get("font", layout.get("font", DEFAULT_FONT)),
        size=size,
        bold=get_block_style(block, layout, role, "bold", False),
    )


def add_field_run(paragraph, field_name):
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = f" {field_name} "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    for element in (begin, instr, separate):
        run = paragraph.add_run()
        run._r.append(element)
    paragraph.add_run("1")
    run = paragraph.add_run()
    run._r.append(end)


def normalize_footer_label(text):
    return "".join(
        chr(ord(ch) + 0xFEE0) if ("0" <= ch <= "9") or ("A" <= ch <= "Z") or ("a" <= ch <= "z") else ch
        for ch in str(text)
    )


def normalize_record_ref(value):
    if not value:
        return ""
    value = str(value).strip()
    if value.lower().startswith("jacar:"):
        value = value.split(":", 1)[1]
    return value


def query_record_ref_by_cache_key(cache_key, db_path=None, *, project_root=None):
    db_path = Path(db_path) if db_path else default_hrs_db_path(project_root)
    if not cache_key or not db_path.exists():
        return ""
    try:
        with sqlite3.connect(str(db_path)) as conn:
            row = conn.execute(
                """
                SELECT native_id, document_id
                FROM document_cache_index
                WHERE cache_key = ?
                   OR cache_basename = ?
                   OR cache_basename = ?
                ORDER BY
                  CASE cache_kind
                    WHEN 'ocr' THEN 0
                    WHEN 'analysis' THEN 1
                    ELSE 2
                  END
                LIMIT 1
                """,
                (cache_key, cache_key, f"{cache_key}.txt"),
            ).fetchone()
    except sqlite3.Error:
        return ""
    if not row:
        return ""
    return normalize_record_ref(row[0] or row[1])


def get_source_cache_key(source):
    record_id = source.get("record_id")
    if record_id:
        return str(record_id)
    ocr_cache = source.get("ocr_cache")
    if ocr_cache:
        return Path(str(ocr_cache)).stem
    return ""


def get_record_ref(source, layout):
    candidates = (
        "jacar_ref",
        "jacar_reference",
        "jacarRef",
        "reference_id",
        "reference",
        "ref",
    )
    for container in (source, layout.get("folio", {}), layout):
        for key in candidates:
            value = container.get(key)
            if value:
                return normalize_record_ref(value)
    return query_record_ref_by_cache_key(get_source_cache_key(source))


def ensure_source_record_ref(source, layout):
    record_ref = get_record_ref(source, layout)
    if record_ref and not source.get("jacar_ref"):
        source["jacar_ref"] = record_ref
    return record_ref



def add_folio_footer_box(section, folio, layout, source_page_label=None, record_ref=None):
    footer = section.footer
    footer.is_linked_to_previous = False
    p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    p.clear()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.space_before = Pt(layout.get("folio", {}).get("space_before_pt", 0))
    pf.space_after = Pt(0)
    font = layout.get("folio", {}).get("font", layout.get("font", DEFAULT_FONT))
    size = layout.get("folio", {}).get("size_pt", 10.5)
    bold = layout.get("folio", {}).get("bold", False)
    footer_parts = []
    if source_page_label:
        footer_parts.append(f"ＨＲＳ：{normalize_footer_label(source_page_label)}")
    else:
        existing_run_count = len(p.runs)
        run = p.add_run("ＨＲＳ：")
        set_run_font(run, font=font, size=size, bold=bold)
        add_field_run(p, "PAGE")
        for run in p.runs[existing_run_count:]:
            set_run_font(run, font=font, size=size, bold=bold)
    if folio:
        footer_parts.append(f"（原史料頁：{normalize_footer_label(folio)}）")
    if record_ref:
        footer_parts.append(f"ＪＡＣＡＲ　Ｒｅｆ．　{normalize_footer_label(record_ref)}")
    if footer_parts:
        run = p.add_run("　".join(footer_parts))
        set_run_font(run, font=font, size=size, bold=bold)


def page_text_length(page):
    return sum(len(block.get("text", "")) for block in page.get("paragraphs", []))


def compact_page_range_label(first_label, last_label):
    if not first_label or first_label == last_label:
        return first_label or last_label
    first_match = re.match(r"^(.*?)(\d+)$", first_label)
    last_match = re.match(r"^(.*?)(\d+)$", last_label)
    if first_match and last_match and first_match.group(1) == last_match.group(1):
        return f"{first_label}~{last_match.group(2)}"
    return f"{first_label}~{last_label}"


def compact_value_range(values):
    values = [str(value) for value in values if value]
    if not values:
        return ""
    first = values[0]
    last = values[-1]
    return first if first == last else f"{first}~{last}"


def group_ocr_pages(pages, max_chars=1100, max_source_pages=3):
    groups = []
    current = []
    current_chars = 0
    for page in pages:
        chars = page_text_length(page)
        would_exceed_chars = current and current_chars + chars > max_chars
        would_exceed_pages = current and len(current) >= max_source_pages
        if would_exceed_chars or would_exceed_pages:
            groups.append(current)
            current = []
            current_chars = 0
        current.append(page)
        current_chars += chars
    if current:
        groups.append(current)
    return groups


def add_page_group(doc, group, layout, record_ref, is_first_group):
    section = doc.sections[0] if is_first_group else doc.add_section(WD_SECTION.NEW_PAGE)
    if not is_first_group:
        apply_section_layout(section, layout)

    first = group[0]
    last = group[-1]
    first_label = first.get("source_page_label") or first.get("page_id")
    last_label = last.get("source_page_label") or last.get("page_id")
    source_label = compact_page_range_label(first_label, last_label)
    folio = compact_value_range(page.get("folio") for page in group)
    add_folio_footer_box(section, folio, layout, source_label, record_ref)

    for page in group:
        for block in page["paragraphs"]:
            add_paragraph_block(doc, block, layout)


def build_doc(json_path_or_schema, out_path, flow_pages=False, merge_pages=True, merge_max_chars=None):
    if isinstance(json_path_or_schema, dict):
        data = json_path_or_schema
    else:
        data = json.loads(Path(json_path_or_schema).read_text(encoding="utf-8"))
    doc = Document()
    layout = data["layout"]
    source = data.get("source", {})
    record_ref = ensure_source_record_ref(source, layout)
    apply_layout(doc, layout)

    if merge_pages and not flow_pages:
        merge_settings = layout.get("ocr_merge", layout.get("translation_merge", {}))
        max_chars = merge_max_chars or merge_settings.get("max_chars_per_page", 1100)
        max_source_pages = merge_settings.get("max_source_pages_per_page", 3)
        groups = group_ocr_pages(data["pages"], max_chars=max_chars, max_source_pages=max_source_pages)
        for group_idx, group in enumerate(groups):
            add_page_group(doc, group, layout, record_ref, group_idx == 0)
        doc.save(out_path)
        return

    for page_idx, page in enumerate(data["pages"]):
        if flow_pages:
            section = doc.sections[0]
            if page_idx == 0:
                source_label = page.get("source_page_label") or page.get("page_id")
                if len(data["pages"]) > 1:
                    last = data["pages"][-1]
                    last_label = last.get("source_page_label") or last.get("page_id")
                    source_label = f"{source_label}-{last_label}"
                add_folio_footer_box(section, page.get("folio"), layout, source_label, record_ref)
        elif page_idx:
            section = doc.add_section(WD_SECTION.NEW_PAGE)
            apply_section_layout(section, layout)
            add_folio_footer_box(
                section,
                page.get("folio"),
                layout,
                page.get("source_page_label") or page.get("page_id"),
                record_ref,
            )
        else:
            section = doc.sections[0]
            add_folio_footer_box(
                section,
                page.get("folio"),
                layout,
                page.get("source_page_label") or page.get("page_id"),
                record_ref,
            )
        for block in page["paragraphs"]:
            add_paragraph_block(doc, block, layout)

    doc.save(out_path)


def parse_args():
    parser = argparse.ArgumentParser(
        description="Generate a Word/WPS-compatible vertical historical-record DOCX from JSON."
    )
    parser.add_argument("input_json", help="Input paragraph-flow JSON file")
    parser.add_argument("output_docx", help="Output .docx file")
    parser.add_argument(
        "--flow-pages",
        action="store_true",
        help=(
            "Do not force each schema pages[] entry to start a new Word page. "
            "Useful when OCR source pages are continuations and should fill available space. "
            "Only the first folio is placed in the footer in this mode."
        ),
    )
    parser.add_argument(
        "--no-merge-pages",
        action="store_true",
        help="Keep one source HRS page per Word page instead of merging OCR pages.",
    )
    parser.add_argument(
        "--merge-max-chars",
        type=int,
        help="Approximate maximum OCR-character count per merged Word page.",
    )
    return parser.parse_args()


if __name__ == "__main__":
    args = parse_args()
    build_doc(
        args.input_json,
        args.output_docx,
        flow_pages=args.flow_pages,
        merge_pages=not args.no_merge_pages,
        merge_max_chars=args.merge_max_chars,
    )
