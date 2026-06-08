"""史料目录导出：DOCX 与 PDF（reportlab）。"""
from __future__ import annotations

from datetime import datetime
from typing import Any

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Emu, Pt
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER
from reportlab.lib.pagesizes import A4, landscape
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import mm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.cidfonts import UnicodeCIDFont
from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

_CATALOG_FONT = "STSong-Light"
_FONT_REGISTERED = False


def _ensure_pdf_font() -> str:
    global _FONT_REGISTERED
    if not _FONT_REGISTERED:
        pdfmetrics.registerFont(UnicodeCIDFont(_CATALOG_FONT))
        _FONT_REGISTERED = True
    return _CATALOG_FONT


# 与 Reports/catalog/史料目录_20260601_202439.docx 对齐的默认版式
_DOCX_HEADERS = ["序号", "专题", "Ref", "标题", "二级分类", "卷名", "馆藏", "页数"]
_DOCX_COL_WIDTHS_EMU = (
    354330,
    516255,
    848360,
    1417955,
    489585,
    3070225,
    1276350,
    393700,
)
_DOCX_BODY_FONT = "Times New Roman"
_DOCX_BODY_FONT_EAST_ASIA = "宋体"
_DOCX_HEADER_FONT = "宋体"
_DOCX_FONT_SIZE = Pt(9)


def _set_landscape_a4(section) -> None:
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    # 与参考文件一致：四边距约 1"/1.25"
    section.left_margin = Emu(914400)
    section.right_margin = Emu(914400)
    section.top_margin = Emu(1143000)
    section.bottom_margin = Emu(1143000)


def _set_cell_vertical_center(cell) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    existing = tc_pr.find(qn("w:vAlign"))
    if existing is not None:
        tc_pr.remove(existing)
    valign = OxmlElement("w:vAlign")
    valign.set(qn("w:val"), "center")
    tc_pr.append(valign)


def _write_table_cell(cell, text: str, *, header: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run(text)
    run.font.size = _DOCX_FONT_SIZE
    if header:
        run.font.name = _DOCX_HEADER_FONT
        run.bold = False
    else:
        run.font.name = _DOCX_BODY_FONT
        r_pr = run._element.get_or_add_rPr()
        r_fonts = r_pr.find(qn("w:rFonts"))
        if r_fonts is None:
            r_fonts = OxmlElement("w:rFonts")
            r_pr.insert(0, r_fonts)
        r_fonts.set(qn("w:eastAsia"), _DOCX_BODY_FONT_EAST_ASIA)
    _set_cell_vertical_center(cell)


def _apply_docx_column_widths(table) -> None:
    for idx, width in enumerate(_DOCX_COL_WIDTHS_EMU):
        table.columns[idx].width = width


def _cell_text(value: Any, *, max_len: int = 120) -> str:
    text = str(value if value is not None else "").strip()
    if len(text) > max_len:
        return text[: max_len - 1] + "…"
    return text


def export_catalog_docx(
    output_path: str,
    *,
    rows: list[dict[str, Any]],
    meta: dict[str, Any],
) -> str:
    del meta  # 参考版式为纯表格，元信息不写入页眉段落
    doc = Document()
    _set_landscape_a4(doc.sections[0])

    headers = _DOCX_HEADERS
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.autofit = False
    _apply_docx_column_widths(table)

    for col, head in enumerate(headers):
        _write_table_cell(table.rows[0].cells[col], head, header=True)

    seq = 0
    for row in rows:
        seq += 1
        values = [
            str(seq),
            _cell_text(row.get("keyword")),
            _cell_text(row.get("ref")),
            _cell_text(row.get("title")),
            _cell_text(row.get("level2")),
            _cell_text(row.get("parent")),
            _cell_text(row.get("repo")),
            _cell_text(row.get("scale")),
        ]
        for col, val in enumerate(values):
            _write_table_cell(table.rows[seq].cells[col], val, header=False)

    doc.save(output_path)
    return output_path


def export_catalog_pdf(
    output_path: str,
    *,
    rows: list[dict[str, Any]],
    meta: dict[str, Any],
) -> str:
    font_name = _ensure_pdf_font()
    page_size = landscape(A4)
    doc = SimpleDocTemplate(
        output_path,
        pagesize=page_size,
        leftMargin=12 * mm,
        rightMargin=12 * mm,
        topMargin=14 * mm,
        bottomMargin=14 * mm,
    )
    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CatTitle",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "CatBody",
        parent=styles["Normal"],
        fontName=font_name,
        fontSize=9,
        leading=12,
    )

    story: list[Any] = []
    story.append(Paragraph("HRS 日本史料目录", title_style))
    story.append(
        Paragraph(
            f"导出时间：{_cell_text(meta.get('exported_at'))}　"
            f"条目数：{meta.get('count', len(rows))}　"
            f"筛选：{_cell_text(meta.get('filter_summary'))}",
            body_style,
        )
    )
    story.append(Spacer(1, 6 * mm))

    headers = _DOCX_HEADERS
    data = [headers]
    for idx, row in enumerate(rows, start=1):
        data.append(
            [
                str(idx),
                _cell_text(row.get("keyword"), max_len=10),
                _cell_text(row.get("ref"), max_len=14),
                _cell_text(row.get("title"), max_len=32),
                _cell_text(row.get("level2"), max_len=10),
                _cell_text(row.get("parent"), max_len=22),
                _cell_text(row.get("repo"), max_len=10),
                _cell_text(row.get("scale")),
            ]
        )

    col_widths = [12 * mm, 18 * mm, 28 * mm, 48 * mm, 16 * mm, 72 * mm, 30 * mm, 12 * mm]
    table = Table(data, colWidths=col_widths, repeatRows=1)
    table.setStyle(
        TableStyle(
            [
                ("FONT", (0, 0), (-1, -1), font_name, 8),
                ("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#d8e4f0")),
                ("TEXTCOLOR", (0, 0), (-1, 0), colors.HexColor("#1a2a3a")),
                ("GRID", (0, 0), (-1, -1), 0.25, colors.grey),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, colors.HexColor("#f6f8fb")]),
            ]
        )
    )
    story.append(table)
    doc.build(story)
    return output_path
