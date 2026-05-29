"""统一 Word (.docx) 导出排版：按字符脚本分 run 设置字体 + 固定行距 + 首行缩进。"""

from __future__ import annotations

import re
from typing import Literal

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor

from utils.text_reflow import looks_like_structured_payload, reflow_page_text

ScriptKind = Literal["latin", "chinese", "japanese"]

# 英文字体（拉丁字母）
FONT_LATIN = "Times New Roman"
# 中文字体
FONT_CHINESE = "宋体"
# 日文字体（必须写入 eastAsia，Word 才对日文汉字/假名生效）
FONT_JAPANESE = "MS Mincho"

BODY_FONT_SIZE_PT = 12
LINE_SPACING_PT = 20
FIRST_LINE_CHARS = "100"  # OOXML：100 = 首行缩进 1 字符

_RE_KANA = re.compile(r"[\u3040-\u309F\u30A0-\u30FF\uFF65-\uFF9F]")
_RE_LATIN = re.compile(r"[A-Za-z0-9]")
_RE_HAN = re.compile(
    r"[\u4E00-\u9FFF\u3400-\u4DBF\uF900-\uFAFF"
    r"\U00020000-\U0002A6DF\U0002F800-\U0002FA1F]"
)
# 互斥标点：避免「、」「。」等共形标点让日文行被误判为中文。
_ZH_PUNCT = set("，；？！""''《》〈〉…—·")
_JA_PUNCT = set("、」「」『』〔〕・〜ー々〻")
# 中日共用标点：不强行改语种，跟随行默认或相邻 run
_SHARED_PUNCT = set("（）。：【】")
# OCR 侧记 / 日式年号等：无假名时也倾向整行日文
_RE_JA_OCR_MARK = re.compile(r"【\s*(?:手写|印)\s*[：:]")
_RE_JA_ARCHIVAL = re.compile(r"昭和|大正|明治|令和|號")


def _line_has_kana(text: str) -> bool:
    return bool(_RE_KANA.search(text))


def _line_default_script(text: str) -> ScriptKind:
    """判定整行默认语种，用于汉字/共用标点的归属。"""
    if _line_has_kana(text):
        return "japanese"
    if _RE_JA_OCR_MARK.search(text) or _RE_JA_ARCHIVAL.search(text):
        return "japanese"
    has_zh = any(ch in _ZH_PUNCT for ch in text)
    has_ja_comma = "、" in text
    if has_zh and not has_ja_comma:
        return "chinese"
    if has_ja_comma and not has_zh:
        return "japanese"
    zh = sum(1 for ch in text if ch in _ZH_PUNCT)
    ja = sum(1 for ch in text if ch in _JA_PUNCT)
    if ja > zh:
        return "japanese"
    if zh > ja:
        return "chinese"
    return "japanese"


def _classify_char(
    ch: str,
    *,
    line_default_script: ScriptKind,
) -> ScriptKind | None:
    if _RE_LATIN.match(ch):
        return "latin"
    if _RE_KANA.match(ch):
        return "japanese"
    if ch in _JA_PUNCT:
        return "japanese"
    if ch in _ZH_PUNCT:
        return "chinese"
    if ch in _SHARED_PUNCT:
        return None
    if _RE_HAN.match(ch):
        return line_default_script
    return None


def _segment_text_by_script(text: str) -> list[tuple[ScriptKind, str]]:
    """将一行文本切成 (script, substring) 片段，供 Word 分 run 设字体。"""
    if not text:
        return []

    line_default_script = _line_default_script(text)

    segments: list[tuple[ScriptKind, str]] = []
    buf: list[str] = []
    current: ScriptKind | None = None

    def flush() -> None:
        nonlocal buf, current
        if not buf:
            return
        kind: ScriptKind = current or line_default_script
        segments.append((kind, "".join(buf)))
        buf = []

    for ch in text:
        kind = _classify_char(ch, line_default_script=line_default_script)
        if kind is None:
            # 空格/符号等：跟随前一个片段，否则按行级默认语种
            if current is not None:
                buf.append(ch)
            else:
                kind = line_default_script
                if current != kind:
                    flush()
                    current = kind
                buf.append(ch)
            continue

        if current != kind:
            flush()
            current = kind
        buf.append(ch)

    flush()
    return segments


def _set_run_fonts(run, script: ScriptKind) -> None:
    """按脚本为 run 设置字体。日文必须写在 eastAsia，不能只写 cs。"""
    run.font.size = Pt(BODY_FONT_SIZE_PT)
    if script == "latin":
        ascii_font = FONT_LATIN
        east_asia = FONT_LATIN
        cs_font = FONT_LATIN
    elif script == "japanese":
        ascii_font = FONT_LATIN
        east_asia = FONT_JAPANESE
        cs_font = FONT_JAPANESE
    else:
        ascii_font = FONT_LATIN
        east_asia = FONT_CHINESE
        cs_font = FONT_CHINESE

    run.font.name = ascii_font
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia)
    r_fonts.set(qn("w:cs"), cs_font)


def _apply_paragraph_style(paragraph) -> None:
    """段前段后 0、固定行距 20 磅、首行缩进 1 字符。"""
    fmt = paragraph.paragraph_format
    fmt.space_before = Pt(0)
    fmt.space_after = Pt(0)
    fmt.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    fmt.line_spacing = Pt(LINE_SPACING_PT)

    p_pr = paragraph._element.get_or_add_pPr()
    ind = p_pr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        p_pr.append(ind)
    ind.set(qn("w:firstLineChars"), FIRST_LINE_CHARS)
    ind.set(qn("w:firstLine"), "0")


def _add_styled_paragraph(doc: Document, text: str):
    paragraph = doc.add_paragraph()
    _apply_paragraph_style(paragraph)
    segments = _segment_text_by_script(text)
    if not segments:
        run = paragraph.add_run("")
        _set_run_fonts(run, "chinese")
        return paragraph

    for script, chunk in segments:
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        _set_run_fonts(run, script)
    return paragraph


def _append_page_break(paragraph) -> None:
    run = paragraph.add_run()
    run.add_break(WD_BREAK.PAGE)


def write_pages_to_docx(
    file_path: str,
    pages: list[str],
    *,
    page_break_between_pages: bool = True,
    reflow_for_readability: bool = True,
) -> None:
    """将多页文本写入 docx。

    reflow_for_readability=True 时，会把 OCR 物理换行合并为自然段（空行仍硬分段），
    避免“一行一段”导致可读性差；JSON 页面自动跳过重排。
    """
    doc = Document()
    usable_pages = [str(p or "").strip() for p in pages if str(p or "").strip()]
    if not usable_pages:
        raise ValueError("导出内容为空")

    for page_idx, page_text in enumerate(usable_pages):
        if reflow_for_readability and not looks_like_structured_payload(page_text):
            paragraph_texts = reflow_page_text(page_text)
        else:
            paragraph_texts = [ln.rstrip("\r") for ln in page_text.splitlines() if ln.strip()]

        if not paragraph_texts:
            paragraph_texts = [""]

        last_paragraph = None
        for content in paragraph_texts:
            last_paragraph = _add_styled_paragraph(doc, content)

        if page_break_between_pages and page_idx < len(usable_pages) - 1 and last_paragraph is not None:
            _append_page_break(last_paragraph)

    doc.save(file_path)


# ------------------------------------------------------------------ #
# 内容1：按页对照报告（横向分页：扫描图 / OCR 文本 / 分析）
# ------------------------------------------------------------------ #

COMPARISON_FONT_OCR = "MS Mincho"
COMPARISON_FONT_ANALYSIS = "宋体"
COMPARISON_HEADER_RGB = RGBColor(0x25, 0x63, 0xEB)
COMPARISON_HEADER_SIZE_PT = 14
COMPARISON_BODY_SIZE_PT = 11
EMU_PER_INCH = 914400

# 对照报告内嵌扫描图：版面按 fit_zoom 适配横向页，光栅化用更高 DPI 再缩放到版面（避免仅 ~72 DPI）
COMPARISON_IMAGE_TARGET_DPI = 240
COMPARISON_IMAGE_MIN_DPI = 96
COMPARISON_IMAGE_MAX_DPI = 600
COMPARISON_IMAGE_MAX_DIMENSION_PX = 10000
PDF_POINTS_PER_INCH = 72.0


def configure_document_landscape(doc: Document) -> None:
    """将文档默认节设为横向 A4。"""
    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width


def landscape_content_size_inches(doc: Document) -> tuple[float, float]:
    """返回当前节可用内容区宽高（英寸）。"""
    section = doc.sections[-1]
    width_emu = section.page_width - section.left_margin - section.right_margin
    height_emu = section.page_height - section.top_margin - section.bottom_margin
    return width_emu / EMU_PER_INCH, height_emu / EMU_PER_INCH


def comparison_image_raster_plan(
    *,
    page_width_pt: float,
    page_height_pt: float,
    content_max_w_in: float,
    content_max_h_in: float,
    target_dpi: float = COMPARISON_IMAGE_TARGET_DPI,
    max_dimension_px: int = COMPARISON_IMAGE_MAX_DIMENSION_PX,
) -> tuple[float, float, float]:
    """计算对照报告扫描图的版面缩放与光栅化缩放。

    返回 (fit_zoom, render_zoom, display_width_inches)。
    - fit_zoom：Word 中的显示尺寸（约 72 DPI 等效版面）
    - render_zoom：PyMuPDF 渲染矩阵，通常 >> fit_zoom，使嵌入 PNG 有效 DPI 接近 target_dpi
    """
    dpi = max(COMPARISON_IMAGE_MIN_DPI, min(COMPARISON_IMAGE_MAX_DPI, float(target_dpi)))
    fit_zoom = min(
        content_max_w_in * PDF_POINTS_PER_INCH / max(page_width_pt, 1.0),
        content_max_h_in * PDF_POINTS_PER_INCH / max(page_height_pt, 1.0),
    )
    fit_zoom = max(0.25, fit_zoom)
    render_zoom = fit_zoom * (dpi / PDF_POINTS_PER_INCH)
    longest_pt = max(page_width_pt, page_height_pt)
    if longest_pt * render_zoom > max_dimension_px:
        render_zoom = max_dimension_px / longest_pt
    display_w_in = min(content_max_w_in, page_width_pt * fit_zoom / PDF_POINTS_PER_INCH)
    return fit_zoom, render_zoom, display_w_in


def add_comparison_page_banner(doc: Document, *, page_index: int, total_pages: int) -> None:
    """页首蓝色「第 n / m 页」横幅。"""
    paragraph = doc.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(10)
    run = paragraph.add_run(f"第 {page_index + 1} / {total_pages} 页")
    run.font.size = Pt(COMPARISON_HEADER_SIZE_PT)
    run.font.bold = True
    run.font.color.rgb = COMPARISON_HEADER_RGB


def append_docx_page_break(doc: Document) -> None:
    doc.add_page_break()


def _set_run_east_asia_font(
    run,
    east_asia_font: str,
    *,
    size_pt: int = COMPARISON_BODY_SIZE_PT,
    bold: bool = False,
) -> None:
    ascii_font = "Times New Roman"
    run.font.name = ascii_font
    run.font.size = Pt(size_pt)
    run.font.bold = bold
    r_pr = run._element.get_or_add_rPr()
    r_fonts = r_pr.find(qn("w:rFonts"))
    if r_fonts is None:
        r_fonts = OxmlElement("w:rFonts")
        r_pr.insert(0, r_fonts)
    r_fonts.set(qn("w:ascii"), ascii_font)
    r_fonts.set(qn("w:hAnsi"), ascii_font)
    r_fonts.set(qn("w:eastAsia"), east_asia_font)
    r_fonts.set(qn("w:cs"), east_asia_font)


def add_comparison_paragraph(doc: Document, text: str, east_asia_font: str) -> None:
    """单段正文，整段使用指定东亚字体（用于 OCR / 分析块）。"""
    paragraph = doc.add_paragraph()
    paragraph.paragraph_format.space_before = Pt(0)
    paragraph.paragraph_format.space_after = Pt(4)
    paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
    paragraph.paragraph_format.line_spacing = Pt(18)
    run = paragraph.add_run(text)
    _set_run_east_asia_font(run, east_asia_font)


def add_comparison_text_body(doc: Document, text: str, east_asia_font: str) -> None:
    """多行正文块。"""
    body = (text or "").strip() or "（空）"
    lines = body.splitlines() or ["（空）"]
    for line in lines:
        add_comparison_paragraph(doc, line if line else " ", east_asia_font)
